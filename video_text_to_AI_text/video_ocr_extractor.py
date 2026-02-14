#!/usr/bin/env python3
"""
Video OCR Extractor - Extract text from video files using Google Cloud Video Intelligence API.

Extracts text from video files (local or GCS) via TEXT_DETECTION (OCR) and outputs
a structured Markdown or DOCX document with deduplication, timestamps, and optional
code formatting.

Authentication: Uses GOOGLE_APPLICATION_CREDENTIALS (Service Account JSON).
"""

from __future__ import annotations

import argparse
import os
import re
import subprocess
import sys
from dataclasses import dataclass
from pathlib import Path

VIDEO_EXTENSIONS = {".mp4", ".mov", ".avi", ".mkv", ".webm", ".m4v"}


def _duration_to_seconds(duration) -> float:
    """Convert protobuf Duration to seconds (handles both .microseconds and .nanos)."""
    seconds = getattr(duration, "seconds", 0) or 0
    nanos = getattr(duration, "nanos", 0) or 0
    micros = getattr(duration, "microseconds", None)
    if micros is not None:
        return seconds + micros * 1e-6
    return seconds + nanos / 1e9


def _format_timestamp(seconds: float) -> str:
    """Format seconds as MM:SS."""
    m = int(seconds // 60)
    s = int(seconds % 60)
    return f"{m:02d}:{s:02d}"


def _looks_like_code(text: str) -> bool:
    """Heuristic: does this text look like Python, JSON, or other code?"""
    stripped = text.strip()
    if not stripped:
        return False
    # Python/JS/JSON indicators
    code_patterns = [
        r"^(def|class|import|from|if|else|elif|for|while|try|except|return|with)\s",
        r"^\s*(if|else|for|while|def|class)\s*\(",
        r"^[\[\{\}\[\]]",  # JSON-like braces
        r"^[\"\'].*[\"\']\s*[:=]",
        r"^\s+#.*$",  # Comment
        r"^\s{2,}[\w]",  # Indented line
        r"=>|=>|&&|\|\||==|!=|===|!==",
        r"^\s*[\w]+\s*\(.*\)\s*:",  # Function/class definition
        r"\.py$|\.json$|\.js$|\.ts$",
        r"^[A-Z_][A-Z0-9_]*\s*=",  # Constants
        r"[\w_]+\.\w+\s*\(",  # Method call: foo.bar(
        r"^\s*[\w_]+\s*=\s*[\w_]+\.",  # var = OBJ.method
    ]
    for pat in code_patterns:
        if re.search(pat, stripped, re.IGNORECASE):
            return True
    # Multiple lines with indentation
    lines = stripped.split("\n")
    if len(lines) > 2 and any(line.startswith("    ") or line.startswith("\t") for line in lines):
        return True
    return False


def _wrap_code_in_backticks(text: str) -> str:
    """Wrap code-like content in triple backticks for Markdown."""
    if not text.strip():
        return text
    if _looks_like_code(text):
        return f"```\n{text}\n```"
    return text


@dataclass
class TextDetection:
    """Single text detection with metadata."""

    text: str
    start_sec: float
    end_sec: float
    confidence: float
    avg_x: float  # For horizontal ordering (left to right)
    avg_y: float  # For vertical ordering (top to bottom, read like image)


def _get_avg_xy(segment) -> tuple[float, float]:
    """Get average (x, y) coordinates from first frame for ordering (read like image)."""
    try:
        if hasattr(segment, "frames") and segment.frames:
            frame = segment.frames[0]
            if hasattr(frame, "rotated_bounding_box") and frame.rotated_bounding_box:
                verts = frame.rotated_bounding_box.vertices
                if verts:
                    avg_x = sum(getattr(v, "x", 0) for v in verts) / len(verts)
                    avg_y = sum(getattr(v, "y", 0) for v in verts) / len(verts)
                    return (avg_x, avg_y)
    except Exception:
        pass
    return (0.0, 0.0)


def _normalize_for_dedup(text: str) -> str:
    """Normalize text for deduplication: collapse whitespace, preserve case for code."""
    return re.sub(r"\s+", " ", text.strip()).strip()


# OCR corruption: mojibake / wrong encoding (e.g. Ã° -> b, Ã© -> e)
_OCR_FIXES = [
    (r"Ã°", "b"),
    (r"Ã©", "e"),
    (r"Ã¨", "e"),
    (r"Ã¡", "a"),
    (r"Ã ", "a"),
    (r"Ã³", "o"),
    (r"Ãº", "u"),
    (r"Ã­", "i"),
    (r"Ã§", "c"),
    (r"Ã±", "n"),
    (r"Ã¶", "o"),
    (r"Ã¼", "u"),
    (r"Ã¢", "a"),
    (r"Ã´", "o"),
    (r"Ã®", "i"),
    (r"Ã»", "u"),
    (r"Å¡", "s"),
    (r"Å¾", "z"),
    (r"ÃŸ", "ss"),
    ("\u2019", "'"),  # right single quote
    ("\u201c", '"'),  # left double quote
    ("\u201d", '"'),  # right double quote
]


def _clean_ocr_text(text: str) -> str:
    """Remove OCR junk: bad UTF, unreadable chars, fix common mojibake."""
    if not text:
        return text
    s = text
    for bad, good in _OCR_FIXES:
        s = s.replace(bad, good)
    # Remove non-printable and control chars (keep ASCII printable + common code chars)
    s = "".join(c for c in s if c.isprintable() or c in "\n\t")
    # Remove other weird unicode (replace with nothing or space)
    s = re.sub(r"[\u200b-\u200f\u2028-\u202f\ufeff]", "", s)
    return s


def _normalize_code_line(text: str) -> str:
    """Normalize code: fix spacing around parens, fix def spacing, collapse interior spaces."""
    s = _clean_ocr_text(text)
    # Collapse multiple spaces
    s = re.sub(r" +", " ", s)
    # Fix: "result ( )" -> "result()", "foo ( bar )" -> "foo(bar)"
    s = re.sub(r"\s*\(\s*\)\s*", "()", s)
    s = re.sub(r"\s*\(\s+", "(", s)
    s = re.sub(r"\s+\)\s*", ")", s)
    # Fix: "def get attribute_value" -> "def get_attribute_value"
    s = re.sub(r"\bdef\s+get\s+attribute_value\b", "def get_attribute_value", s, flags=re.IGNORECASE)
    s = re.sub(r"\bdef\s+get\s+attribute_date\b", "def get_attribute_date", s, flags=re.IGNORECASE)
    s = re.sub(r"\bdef\s+get\s+table_names\b", "def get_table_names", s, flags=re.IGNORECASE)
    # Fix: "attritute" -> "attribute" (common OCR typo)
    s = re.sub(r"\battritute\b", "attribute", s, flags=re.IGNORECASE)
    # Fix: "attribute value" in identifiers -> "attribute_value"
    s = re.sub(r"attribute\s+value", "attribute_value", s, flags=re.IGNORECASE)
    s = re.sub(r"attribute\s+date", "attribute_date", s, flags=re.IGNORECASE)
    s = re.sub(r"query\s+job", "query_job", s, flags=re.IGNORECASE)
    s = re.sub(r"curr\s+table", "curr_table", s, flags=re.IGNORECASE)
    s = re.sub(r"subject\s+area", "subject_area", s, flags=re.IGNORECASE)
    s = re.sub(r"table\s+name", "table_name", s, flags=re.IGNORECASE)
    s = re.sub(r"attribute\s+_", "attribute_", s)
    s = re.sub(r"_\s+value", "_value", s)
    s = re.sub(r"que\s+ry\s+job", "query_job", s)
    s = re.sub(r"column\s+\.", "column.", s)
    s = re.sub(r"\.\s+result", ".result", s)
    s = re.sub(r"curr_table_name\s+\.\s+", "curr_table_name.", s)
    s = re.sub(r"PROJECT\s+ID", "PROJECT_ID", s, flags=re.IGNORECASE)
    s = re.sub(r"column\s+name", "column_name", s, flags=re.IGNORECASE)
    s = re.sub(r"\.\s+lowe\s*r\s*\.", ".lower().", s)
    s = re.sub(r"=>", "->", s)
    s = re.sub(r"query_job\s*,\s*result", "query_job.result", s)
    s = re.sub(r"def\s+get\s+closest_column", "def get_closest_column", s)
    s = re.sub(r"def\s+get\s+table_data", "def get_table_data", s)
    s = re.sub(r"gettable_data", "get_table_data", s)
    s = re.sub(r"getattribute_date", "get_attribute_date", s)
    s = re.sub(r"getattribute_value", "get_attribute_value", s)
    return s.strip()


def _is_ui_junk(line: str) -> bool:
    """Filter editor UI, status bar, file explorer, truncated fragments, junk."""
    s = line.strip().lower()
    if not s or len(s) < 4:
        return True
    if re.search(r"\b(public|private|void|class|interface|import|package|return|static)\s", s):
        return False
    if s.endswith("..."):
        return True
    if re.match(r"^(ate_|eneric_|=_|ble_|unction_|chema_|eric_|ic_table|etadata_|enerate_test|enerate_schema)[a-z_]*(\.py)?\.?$", s):
        return True
    if re.match(r"^(est_model|bt_model|chema_yml|unction_generator|t_model_files)[a-z_]*\.py$", s):
        return True
    if re.match(r"^(unc_gen|nc_gen|etadata_csv|nodel_files)(\.py)?$", s):
        return True
    if re.match(r"^(_data|_csv)\.py$", s) and len(s) < 15:
        return True
    if re.match(r"^es\.py$", s):
        return True
    if s in ("_data", "tions", "ations", "ns.py", "ons.py"):
        return True
    if re.search(r"\bln\s*\d+.*col\s*\d+|\(\d+\s+selected\)|spaces:\s*\d|utf-8\s*lf", s):
        return True
    if re.search(r"\.py\s*\d+\s*x\s*$|,\s*py\s+\d+\s+x", s):
        return True
    if re.match(r"^[a-z]\s+[a-z_][a-z0-9_]*$", s) and len(s) < 25:
        return True
    if s in ("bute_value", "esis", "lests", "lesls", "tesis", "explorer"):
        return True
    if s == "explorer" or s.startswith("opt_imp_") or "pt_imp_scripts" in s:
        return True
    if re.match(r"^v\s+(src|imp|dbt|utils|generators|operations|table_data)\b", s):
        return True
    if re.match(r"^[a-z]{1,3}_[a-z]+\.py\s*$", s):
        return True
    if ("/" in s or ">" in s) and len(s) < 50:
        return True
    if re.search(r"imp\s*code\s*>|>.*>.*>", s) and "def " not in s and "import " not in s:
        return True
    if re.match(r"^[a-z]{1,4}\.(yml|yaml|md|json|db)\s*$", s):
        return True
    if re.match(r"^[a-z]+_[a-z]+_[a-z]+\.py\.\.$", s):
        return True
    if re.match(r"^(unctions|ble_|e_|ct\.|e_data|le_data|chema|_gen|ons|nc_gen|ema)\.?py?$", s):
        return True
    if s in ("arfarme.md", "jenkinsfile", "gitignore", "imp_code"):
        return True
    if re.match(r"^[a-z]+-clinical-store", s):
        return True
    if re.match(r"^(pt\s+imp\s+scripts|vai-clinical|bq\s+sql_func_gen|uu\s+impui)", s):
        return True
    if "unctions generators" in s or "get table_data.py" in s:
        return True
    if re.match(r"^[a-z]_(csv|schema|test|drop|dbt|function)\w*\.py$", s):
        return True
    if re.match(r"^e_[a-z_]+\.py$", s) and len(s) < 30:
        return True
    if "(pythom" in s or "lf(python" in s or "d.bigquery:" in s and "import" in s:
        return True
    if re.match(r"^\(function\)\s+def", s):
        return True
    if re.match(r"^[a-z]\.[a-z]+:\s*table import", s):
        return True
    return False


def _code_canonical(text: str) -> str:
    """Canonical form for fuzzy dedup: alphanumeric + underscore only, lowercase."""
    s = _normalize_code_line(text)
    s = re.sub(r"[^a-zA-Z0-9_]", "", s).lower()
    return s


def _code_quality_score(text: str) -> float:
    """Higher = more likely valid code (Python or Java). Penalize junk, reward valid patterns."""
    s = text.strip()
    if len(s) < 4:
        return 0
    score = 0
    weird = len([c for c in s if ord(c) > 127 or (not c.isalnum() and c not in " _()[]{}:=.,'\"/\\\t\n#-;<>")])
    score -= weird * 2
    # Python patterns
    if re.search(r"\bdef\s+\w+\s*\(", s):
        score += 3
    if "->" in s:
        score += 2
    # Java patterns
    if re.search(r"\b(public|private|protected|void|class|interface|static|final)\s", s):
        score += 3
    if "{" in s or "}" in s:
        score += 2
    if ";" in s or s.endswith(";"):
        score += 1
    if "=" in s and ("(" in s or ";" in s):
        score += 2
    if re.match(r"^[a-zA-Z_][a-zA-Z0-9_]*\s*=\s*", s):
        score += 1
    if re.search(r"\b(import|package|return|if|for|while)\s", s):
        score += 2
    if re.match(r"^[0-9]+$", s):
        score -= 5
    if len(s.split()) == 1 and len(s) < 5:
        score -= 1
    return score


def detect_text(
    input_path: str,
    credentials_path: str | None = None,
    language_hints: list[str] | None = None,
    min_duration_sec: float = 0.0,
):
    """
    Call Google Cloud Video Intelligence API for TEXT_DETECTION.

    Args:
        input_path: Local file path or gs:// GCS URI.
        credentials_path: Path to service account JSON. Uses GOOGLE_APPLICATION_CREDENTIALS if None.
        language_hints: Optional language hints, e.g. ["en-US"].
        min_duration_sec: Ignore text visible for less than this (reduces noise).

    Returns:
        AnnotateVideoResponse from the API.
    """
    try:
        from google.cloud import videointelligence_v1 as videointelligence
    except ImportError:
        print(
            "Error: google-cloud-videointelligence is required. Run:\n"
            "  pip install google-cloud-videointelligence",
            file=sys.stderr,
        )
        sys.exit(1)

    if credentials_path:
        os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = credentials_path
    elif not os.environ.get("GOOGLE_APPLICATION_CREDENTIALS"):
        # Use Application Default Credentials (from: gcloud auth application-default login)
        adc_path = Path.home() / ".config" / "gcloud" / "application_default_credentials.json"
        if adc_path.exists():
            pass  # Client will auto-discover ADC
        else:
            print(
                "Note: No credentials set. Run one of:\n"
                "  1. gcloud auth application-default login\n"
                "  2. export GOOGLE_APPLICATION_CREDENTIALS=/path/to/service-account.json",
                file=sys.stderr,
            )

    # Use project for billing when using ADC (gcloud auth application-default login)
    client_options = None
    project = os.environ.get("GOOGLE_CLOUD_PROJECT")
    if project:
        from google.api_core.client_options import ClientOptions
        client_options = ClientOptions(quota_project_id=project)

    video_client = videointelligence.VideoIntelligenceServiceClient(client_options=client_options)

    features = [videointelligence.Feature.TEXT_DETECTION]
    video_context = videointelligence.VideoContext()
    if language_hints:
        video_context.text_detection_config = videointelligence.TextDetectionConfig(
            language_hints=language_hints
        )

    is_gcs = input_path.strip().lower().startswith("gs://")

    try:
        if is_gcs:
            request = {
                "features": features,
                "input_uri": input_path,
                "video_context": video_context,
            }
            print(f"Processing GCS video: {input_path}", file=sys.stderr)
        else:
            path = Path(input_path)
            if not path.exists():
                raise FileNotFoundError(f"Video file not found: {input_path}")
            with open(path, "rb") as f:
                input_content = f.read()
            request = {
                "features": features,
                "input_content": input_content,
                "video_context": video_context,
            }
            print(f"Processing local video: {input_path}", file=sys.stderr)

        operation = video_client.annotate_video(request=request)
        print("Waiting for long-running operation to complete (may take several minutes)...", file=sys.stderr)
        result = operation.result(timeout=600)
        return result

    except Exception as e:
        if "quota" in str(e).lower() or "429" in str(e):
            print("Error: API quota exceeded. Try again later or increase quota.", file=sys.stderr)
        elif "403" in str(e) or "permission" in str(e).lower():
            print(
                "Error: Authentication/permission failed. Check your service account and API enablement.",
                file=sys.stderr,
            )
        elif "404" in str(e):
            print("Error: Video file or GCS object not found.", file=sys.stderr)
        raise


def process_results(
    result,
    min_duration_sec: float = 0.0,
    code_mode: bool = False,
) -> list[TextDetection]:
    """
    Parse annotation_results, clean OCR, deduplicate (including fuzzy), sort.
    code_mode: apply code normalization and fuzzy dedup for similar lines.
    """
    raw: list[TextDetection] = []

    if not result.annotation_results:
        return []

    for ann_result in result.annotation_results:
        text_annotations = getattr(ann_result, "text_annotations", []) or []
        for ta in text_annotations:
            text = (ta.text or "").strip()
            if not text:
                continue

            segments = getattr(ta, "segments", []) or []
            for seg in segments:
                video_seg = getattr(seg, "segment", None)
                if not video_seg:
                    continue
                start = _duration_to_seconds(getattr(video_seg, "start_time_offset", None) or 0)
                end = _duration_to_seconds(getattr(video_seg, "end_time_offset", None) or 0)
                duration = end - start

                if duration < min_duration_sec:
                    continue

                confidence = getattr(seg, "confidence", 1.0) or 1.0
                if confidence < 0.5:
                    continue

                avg_x, avg_y = _get_avg_xy(seg)
                raw.append(
                    TextDetection(text=text, start_sec=start, end_sec=end, confidence=confidence, avg_x=avg_x, avg_y=avg_y)
                )

    # Clean and optionally normalize code
    cleaned: list[TextDetection] = []
    for d in raw:
        t = _clean_ocr_text(d.text)
        if code_mode:
            t = _normalize_code_line(t)
        t = t.strip()
        if len(t) < 3:
            continue
        if re.match(r"^[\s\-_=\.]+$", t):
            continue
        if code_mode and _is_ui_junk(t):
            continue
        cleaned.append(TextDetection(text=t, start_sec=d.start_sec, end_sec=d.end_sec, confidence=d.confidence, avg_x=d.avg_x, avg_y=d.avg_y))

    if code_mode and not cleaned and raw:
        for d in raw:
            t = _clean_ocr_text(d.text)
            t = t.strip()
            if len(t) >= 3 and not re.match(r"^[\s\-_=\.]+$", t):
                cleaned.append(TextDetection(text=t, start_sec=d.start_sec, end_sec=d.end_sec, confidence=d.confidence, avg_x=d.avg_x, avg_y=d.avg_y))

    # Deduplicate: exact match first
    seen_exact: dict[str, TextDetection] = {}
    for d in cleaned:
        key = _normalize_for_dedup(d.text)
        if key not in seen_exact or d.confidence > seen_exact[key].confidence:
            seen_exact[key] = d

    if not code_mode:
        detections = list(seen_exact.values())
        detections.sort(key=lambda x: (x.start_sec, x.avg_y, x.avg_x))
        return detections

    # Fuzzy dedup for code: group by canonical form, pick best quality
    canonical_to_best: dict[str, TextDetection] = {}
    for d in seen_exact.values():
        canon = _code_canonical(d.text)
        if len(canon) < 3:
            continue
        q = _code_quality_score(d.text)
        if canon not in canonical_to_best or q > _code_quality_score(canonical_to_best[canon].text):
            canonical_to_best[canon] = d

    detections = list(canonical_to_best.values())
    if not detections and seen_exact:
        detections = list(seen_exact.values())
    detections.sort(key=lambda x: (x.start_sec, x.avg_y, x.avg_x))
    return detections


def _format_code_with_indentation(detections: list[TextDetection]) -> str:
    """
    Format code with semantic indentation and blank lines between defs.
    - def/class/from/import at column 0
    - Function/loop bodies at 4 spaces
    - Nested blocks at 8 spaces
    """
    if not detections:
        return ""
    SPACES = 4
    output_lines: list[str] = []
    current_indent = 0
    prev_was_def_or_class = False

    for d in detections:
        text = d.text.strip()
        if not text:
            continue

        is_def_or_class = bool(re.match(r"^(def|class)\s+\w+", text))
        is_top_level = bool(re.match(r"^(from|import|@)\s", text)) or re.match(r"^(CLIENT|PROJECT_ID)\s*=", text)
        is_loop_or_cond = bool(re.match(r"^(for|if|elif|else|while|try|except|with)\s", text))
        is_return = bool(re.match(r"^return\s", text))

        if is_def_or_class or is_top_level:
            indent = 0
            current_indent = 0
        elif is_loop_or_cond:
            indent = current_indent if current_indent > 0 else SPACES
            if current_indent == 0:
                current_indent = SPACES
        elif is_return or (current_indent > 0 and re.search(r"^\w+\s*=|\.append|\.result|\.query", text)):
            indent = current_indent
        elif current_indent > 0:
            indent = current_indent
        else:
            indent = 0

        indent_str = " " * indent
        if output_lines and is_def_or_class and not prev_was_def_or_class:
            output_lines.append("")
        prev_was_def_or_class = is_def_or_class

        output_lines.append(indent_str + text)

    return "\n".join(output_lines)


def generate_markdown(
    detections: list[TextDetection],
    output_path: str,
    wrap_code: bool = True,
    no_timestamps: bool = False,
    code_mode: bool = False,
) -> None:
    """
    Write detections to a Markdown file.
    code_mode: no timestamps, indentation from position, blank lines between defs.
    """
    if code_mode:
        lines = ["# Extracted Code", ""]
        code_block = _format_code_with_indentation(detections)
        if not code_block.strip():
            code_block = "# No text detected in video.\n# Possible causes: video very short, low resolution, fast scrolling, or format/codec issues.\n# Try: re-record with slower scrolling, higher resolution, or convert to MP4."
        lines.append("```python")
        lines.append(code_block)
        lines.append("```")
    else:
        lines = [
            "# Video OCR Extracted Text",
            "",
            "Text extracted from video using Google Cloud Video Intelligence API (TEXT_DETECTION).",
            "",
            "---",
            "",
        ]
        for d in detections:
            content = d.text
            if wrap_code:
                content = _wrap_code_in_backticks(content)
            if no_timestamps:
                lines.append(content)
                lines.append("")
            else:
                ts_start = _format_timestamp(d.start_sec)
                ts_end = _format_timestamp(d.end_sec)
                lines.append(f"**[{ts_start} - {ts_end}]**")
                lines.append("")
                lines.append(content)
                lines.append("")

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    print(f"Wrote Markdown to {output_path}", file=sys.stderr)


def generate_docx(
    detections: list[TextDetection],
    output_path: str,
    wrap_code: bool = True,
    no_timestamps: bool = False,
    code_mode: bool = False,
) -> None:
    """Write detections to a DOCX file using python-docx."""
    try:
        from docx import Document
    except ImportError:
        print(
            "Error: python-docx is required for .docx output. Run:\n"
            "  pip install python-docx",
            file=sys.stderr,
        )
        sys.exit(1)

    doc = Document()
    doc.add_heading("Extracted Code" if code_mode else "Video OCR Extracted Text", 0)
    if not code_mode:
        doc.add_paragraph(
            "Text extracted from video using Google Cloud Video Intelligence API (TEXT_DETECTION)."
        )
        doc.add_paragraph()

    for d in detections:
        if not code_mode and not no_timestamps:
            ts_start = _format_timestamp(d.start_sec)
            ts_end = _format_timestamp(d.end_sec)
            doc.add_paragraph(f"[{ts_start} - {ts_end}]", style="Heading 2")
        doc.add_paragraph(d.text, style="Normal")

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"Wrote DOCX to {output_path}", file=sys.stderr)


def list_videos_in_folder(folder_path: str) -> list[str]:
    """List video file paths in a folder. Supports local paths and gs:// GCS URIs."""
    paths: list[str] = []
    folder = folder_path.rstrip("/")

    if folder.lower().startswith("gs://"):
        try:
            result = subprocess.run(
                ["gsutil", "ls", f"{folder}/"],
                capture_output=True,
                text=True,
                timeout=60,
            )
            if result.returncode == 0 and result.stdout:
                for line in result.stdout.strip().split("\n"):
                    line = line.strip()
                    if line and not line.endswith("/"):
                        ext = Path(line).suffix.lower()
                        if ext in VIDEO_EXTENSIONS:
                            paths.append(line)
        except (subprocess.CalledProcessError, FileNotFoundError) as e:
            print(f"Error listing GCS folder: {e}", file=sys.stderr)
    else:
        p = Path(folder)
        if not p.exists() or not p.is_dir():
            print(f"Error: Input folder not found: {folder}", file=sys.stderr)
            return []
        for f in p.iterdir():
            if f.is_file() and f.suffix.lower() in VIDEO_EXTENSIONS:
                paths.append(str(f.absolute()))

    return sorted(paths)


def process_single_video(
    input_path: str,
    output_path: str,
    credentials_path: str | None,
    lang_hints: list[str] | None,
    min_dur: float,
    code_mode: bool,
    wrap_code: bool,
    no_timestamps: bool,
    verbose: bool = False,
) -> bool:
    """Process one video and write output. Returns True on success."""
    try:
        result = detect_text(
            input_path=input_path,
            credentials_path=credentials_path,
            language_hints=lang_hints,
            min_duration_sec=min_dur,
        )
        if verbose and result.annotation_results:
            total = sum(len(getattr(ar, "text_annotations", []) or []) for ar in result.annotation_results)
            print(f"  Raw text annotations from API: {total}", file=sys.stderr)
        detections = process_results(result, min_duration_sec=min_dur, code_mode=code_mode)
        if verbose:
            print(f"  Detections after processing: {len(detections)}", file=sys.stderr)
        if not detections and result.annotation_results:
            ann = result.annotation_results[0]
            raw_texts = getattr(ann, "text_annotations", []) or []
            if raw_texts and verbose:
                print(f"  Sample raw text: {raw_texts[0].text[:80]!r}...", file=sys.stderr)
        if output_path.lower().endswith(".docx"):
            generate_docx(detections, output_path, wrap_code=wrap_code, no_timestamps=no_timestamps, code_mode=code_mode)
        else:
            generate_markdown(detections, output_path, wrap_code=wrap_code, no_timestamps=no_timestamps, code_mode=code_mode)
        return True
    except Exception as e:
        print(f"Error processing {input_path}: {e}", file=sys.stderr)
        return False


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Extract text from video using Google Cloud Video Intelligence API (OCR)."
    )
    parser.add_argument(
        "--input", "-i",
        help="Path to video file (local) or gs://bucket/path (GCS URI). Required unless --input-folder is used.",
    )
    parser.add_argument(
        "--output", "-o",
        help="Output file path (.md or .docx). Required for single-file mode.",
    )
    parser.add_argument(
        "--input-folder",
        help="Process all videos in this folder (local or gs://). Outputs go to --output-folder.",
    )
    parser.add_argument(
        "--output-folder",
        help="Output folder for batch mode. Writes <video_name>_output.md for each input.",
    )
    parser.add_argument(
        "--credentials", "-c",
        default=None,
        help="Path to service account JSON. Default: GOOGLE_APPLICATION_CREDENTIALS env var.",
    )
    parser.add_argument(
        "--language",
        default="en-US",
        help="Language hint for OCR (e.g. en-US, es, ja). Default: en-US.",
    )
    parser.add_argument(
        "--min-duration",
        type=float,
        default=0.0,
        help="Ignore text visible for less than N seconds (reduces noise). Default: 0.",
    )
    parser.add_argument(
        "--no-code-formatting",
        action="store_true",
        help="Do not wrap code-like text in triple backticks.",
    )
    parser.add_argument(
        "--code-mode",
        action="store_true",
        help="Clean code output: no timestamps, no duplicate lines, ordered like reading an image.",
    )
    parser.add_argument(
        "--no-timestamps",
        action="store_true",
        help="Omit timestamps from output (use with --code-mode for code videos).",
    )
    parser.add_argument(
        "--verbose", "-v",
        action="store_true",
        help="Print debug info (raw detection count, etc.).",
    )
    args = parser.parse_args()

    lang_hints = [args.language] if args.language else None
    min_dur = max(args.min_duration, 0.2) if args.code_mode else args.min_duration
    code_mode = args.code_mode
    wrap_code = not args.no_code_formatting
    no_timestamps = args.no_timestamps or code_mode

    if args.input_folder and args.output_folder:
        # Batch mode: process all videos in input folder
        videos = list_videos_in_folder(args.input_folder)
        if not videos:
            print("No video files found in input folder.", file=sys.stderr)
            sys.exit(1)
        print(f"Found {len(videos)} video(s) to process.", file=sys.stderr)
        out_folder = args.output_folder.rstrip("/")
        is_gcs_output = out_folder.lower().startswith("gs://")
        if not is_gcs_output:
            Path(out_folder).mkdir(parents=True, exist_ok=True)
        success = 0
        for i, input_path in enumerate(videos):
            base = Path(input_path).stem
            if is_gcs_output:
                output_path = str(Path.cwd() / f".tmp_{base}_output.md")
            else:
                output_path = str(Path(out_folder) / f"{base}_output.md")
            print(f"[{i+1}/{len(videos)}] Processing: {input_path}", file=sys.stderr)
            if process_single_video(
                input_path, output_path,
                args.credentials, lang_hints, min_dur,
                code_mode, wrap_code, no_timestamps,
                verbose=args.verbose,
            ):
                success += 1
                if is_gcs_output:
                    try:
                        subprocess.run(
                            ["gsutil", "cp", output_path, f"{out_folder}/{base}_output.md"],
                            check=True, capture_output=True, timeout=60,
                        )
                        print(f"  Uploaded to {out_folder}/{base}_output.md", file=sys.stderr)
                    except Exception as e:
                        print(f"  Upload failed: {e}", file=sys.stderr)
                    try:
                        os.remove(output_path)
                    except OSError:
                        pass
            else:
                print(f"  Skipped (error).", file=sys.stderr)
        print(f"Completed: {success}/{len(videos)} videos processed.", file=sys.stderr)
        return

    if not args.input or not args.output:
        parser.error("--input and --output are required (or use --input-folder and --output-folder for batch).")

    input_path = args.input
    output_path = args.output
    ext = Path(output_path).suffix.lower()
    if ext not in (".md", ".docx"):
        if ext == "":
            output_path = str(Path(output_path).with_suffix(".md"))

    process_single_video(
        input_path, output_path,
        args.credentials, lang_hints, min_dur,
        code_mode, wrap_code, no_timestamps,
        verbose=args.verbose,
    )


if __name__ == "__main__":
    main()
